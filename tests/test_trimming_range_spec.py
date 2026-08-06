"""
Unit and Integration Tests for Document Trimming with Page Range Specifications
"""
import os
import pytest
import docx

from word_exporter_pro.core.com_engine import PageExporterEngine
from word_exporter_pro.core.range_parser import PageRangeParser
from word_exporter_pro.core.batch_processor import BatchProcessor, ExportJobConfig


@pytest.fixture
def sample_multi_page_docx(tmp_path):
    """Creates a 3-page DOCX file with distinct paragraph content on each page."""
    doc_path = os.path.join(str(tmp_path), "MultiPageSample.docx")
    doc = docx.Document()

    # Page 1
    doc.add_paragraph("Paragraph 1 on Page 1")
    doc.add_paragraph("Paragraph 2 on Page 1")
    p1_break = doc.add_paragraph()
    p1_break.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

    # Page 2
    doc.add_paragraph("Paragraph 1 on Page 2")
    doc.add_paragraph("Paragraph 2 on Page 2")
    p2_break = doc.add_paragraph()
    p2_break.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

    # Page 3
    doc.add_paragraph("Paragraph 1 on Page 3")
    doc.add_paragraph("Paragraph 2 on Page 3")

    doc.save(doc_path)
    return doc_path


def test_fallback_docx_trimming_page_1(sample_multi_page_docx, tmp_path):
    """Verifies that python-docx fallback engine exports Page 1 content cleanly without dropping it."""
    out_file = os.path.join(str(tmp_path), "Exported_Page1.docx")
    
    PageExporterEngine._export_by_docx_fallback(
        abs_source=sample_multi_page_docx,
        abs_output=out_file,
        start_page=1,
        end_page=1,
        export_format="docx",
        total_pages=3
    )

    assert os.path.exists(out_file)
    trimmed_doc = docx.Document(out_file)
    full_text = "\n".join(p.text for p in trimmed_doc.paragraphs if p.text.strip())
    
    assert "Page 1" in full_text
    assert "Page 2" not in full_text
    assert "Page 3" not in full_text


def test_fallback_docx_trimming_page_2(sample_multi_page_docx, tmp_path):
    """Verifies that python-docx fallback engine exports Page 2 content correctly."""
    out_file = os.path.join(str(tmp_path), "Exported_Page2.docx")
    
    PageExporterEngine._export_by_docx_fallback(
        abs_source=sample_multi_page_docx,
        abs_output=out_file,
        start_page=2,
        end_page=2,
        export_format="docx",
        total_pages=3
    )

    assert os.path.exists(out_file)
    trimmed_doc = docx.Document(out_file)
    full_text = "\n".join(p.text for p in trimmed_doc.paragraphs if p.text.strip())
    
    assert "Page 2" in full_text
    assert "Page 1" not in full_text
    assert "Page 3" not in full_text


def test_batch_processor_all_individual_split(sample_multi_page_docx, tmp_path):
    """Verifies batch processor splitting document with 'all-individual' range spec into separate files."""
    output_dir = os.path.join(str(tmp_path), "out_split")
    config = ExportJobConfig(
        source_files=[sample_multi_page_docx],
        range_expression="all-individual",
        output_dir=output_dir,
        export_format="docx",
        engine_mode="trimming",
        overwrite=True
    )

    processor = BatchProcessor(config)
    created_files = []

    def on_file_created(path):
        created_files.append(path)

    # Run batch export synchronously
    processor._run_job(on_progress=None, on_finished=None, on_file_created=on_file_created)

    assert len(created_files) == 3
    for fpath in created_files:
        assert os.path.exists(fpath)
        assert os.path.getsize(fpath) > 0


def test_custom_multi_range_split(sample_multi_page_docx, tmp_path):
    """Verifies that multi-range spec '1-2, 3' splits document into separate distinct range files."""
    output_dir = os.path.join(str(tmp_path), "out_multi_range")
    config = ExportJobConfig(
        source_files=[sample_multi_page_docx],
        range_expression="1-2, 3",
        output_dir=output_dir,
        export_format="docx",
        engine_mode="trimming",
        overwrite=True
    )

    processor = BatchProcessor(config)
    created_files = []

    def on_file_created(path):
        created_files.append(path)

    processor._run_job(on_progress=None, on_finished=None, on_file_created=on_file_created)

    assert len(created_files) == 2
    
    # File 1 (Pages 1-2)
    doc1 = docx.Document(created_files[0])
    text1 = "\n".join(p.text for p in doc1.paragraphs if p.text.strip())
    assert "Page 1" in text1 or "Page 2" in text1
    assert "Page 3" not in text1

    # File 2 (Page 3)
    doc2 = docx.Document(created_files[1])
    text2 = "\n".join(p.text for p in doc2.paragraphs if p.text.strip())
    assert "Page 3" in text2
    assert "Page 1" not in text2
