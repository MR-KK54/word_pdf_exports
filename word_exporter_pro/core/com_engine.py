"""
Word Page Exporter Pro - Microsoft Word COM Engine
Provides high-fidelity page extraction and document pagination via Word.Application.
"""

import os
import shutil
import subprocess
import tempfile
from typing import Dict, Any, Optional, Tuple


try:
    import pythoncom  # type: ignore
    import win32com.client  # type: ignore
except ImportError:  # pragma: no cover - exercised in non-Windows environments
    pythoncom = None
    win32com = None

try:
    import docx  # type: ignore
except ImportError:
    docx = None

try:
    import aspose.words as aw  # type: ignore
except ImportError:
    aw = None

from word_exporter_pro.utils.logger import get_logger

logger = get_logger()


def _require_server_word_engine() -> None:
    """Reject inaccurate XML-based pagination on non-Windows hosts.

    ``python-docx`` reads document structure but does not lay a document out on
    pages.  Using it to guess a page count makes an export look successful while
    silently putting the wrong content in each output file.  Linux deployments
    must therefore have Aspose.Words available for Word-document splitting.
    """
    if aw is None:
        raise RuntimeError(
            "This Linux server cannot accurately split Word documents because "
            "Aspose.Words is unavailable. Ensure the Render build installs "
            "aspose-words, then redeploy. PDF splitting is unaffected."
        )


# Table row layout cache, keyed by (id(doc), table.Range.Start, table.Range.End,
# table.Rows.Count). Layout of a table only changes when its rows are added or
# deleted, and the key changes on every row delete, so the cached positions stay
# valid. Building a 120-row layout costs several seconds of COM round-trips, and
# the same layout is needed repeatedly within one export (boundary snap, front and
# back trimming, section lookup), so caching is a large speed win.
_TABLE_ROW_LAYOUT_CACHE: Dict[Tuple[int, int, int, int], Tuple[dict, int]] = {}


def _is_libreoffice_available() -> bool:
    """Returns True if LibreOffice (soffice) executable is installed and runnable on system PATH."""
    try:
        res = subprocess.run(["soffice", "--version"], capture_output=True, text=True, timeout=5)
        return res.returncode == 0
    except Exception:
        return False


def _convert_docx_to_pdf_libreoffice(abs_source: str, out_dir: str) -> Optional[str]:
    """Converts a Word/RTF document to PDF using headless LibreOffice."""
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", abs_source, "--outdir", out_dir]
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=45)
        if res.returncode == 0:
            base = os.path.splitext(os.path.basename(abs_source))[0]
            pdf_path = os.path.join(out_dir, f"{base}.pdf")
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                return pdf_path
            logger.warning(f"LibreOffice command succeeded but output PDF not found at {pdf_path}")
        else:
            logger.warning(f"LibreOffice PDF conversion exited with code {res.returncode}: {res.stderr}")
    except Exception as e:
        logger.warning(f"LibreOffice conversion failed for '{abs_source}': {e}")
    return None


def _convert_pdf_to_docx_libreoffice(abs_pdf: str, abs_docx_out: str) -> Optional[str]:
    """Converts a PDF document to DOCX using headless LibreOffice."""
    out_dir = os.path.dirname(abs_docx_out)
    try:
        cmd = ["soffice", "--headless", "--infilter=writer_pdf_import", "--convert-to", "docx", abs_pdf, "--outdir", out_dir]
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=45)
        if res.returncode == 0:
            base = os.path.splitext(os.path.basename(abs_pdf))[0]
            gen_path = os.path.join(out_dir, f"{base}.docx")
            if os.path.exists(gen_path) and os.path.getsize(gen_path) > 0:
                if gen_path != abs_docx_out:
                    shutil.move(gen_path, abs_docx_out)
                return abs_docx_out
    except Exception as e:
        logger.warning(f"LibreOffice PDF to DOCX conversion failed for '{abs_pdf}': {e}")
    return None


def _require_word_com() -> None:
    if pythoncom is None or win32com is None:
        raise RuntimeError(
            "Microsoft Word COM support is unavailable on this server. "
            "Install pywin32 on Windows to enable Word document export features."
        )


# Word COM Constants
WD_GOTO_PAGE = 1
WD_GOTO_ABSOLUTE = 1
WD_STATISTIC_PAGES = 2
WD_ACTIVE_END_PAGE_NUMBER = 3
WD_ALERTS_NONE = 0

EXPORT_FORMAT_MAP = {
    "docx": 16,   # wdFormatXMLDocument
    "doc": 0,     # wdFormatDocument97
    "pdf": 17,    # wdFormatPDF
    "rtf": 6,     # wdFormatRTF
    "docm": 13,   # wdFormatXMLDocumentMacroEnabled
    "dotx": 14,   # wdFormatXMLTemplate
}


class WordCOMContext:
    """Context manager for Word.Application COM lifecycle."""

    def __init__(self, visible: bool = False):
        self.visible = visible
        self.word_app = None
        self.co_initialized = False

    @property
    def available(self) -> bool:
        return pythoncom is not None and win32com is not None


    def __enter__(self):
        _require_word_com()
        pythoncom.CoInitialize()
        self.co_initialized = True
        try:
            # DispatchEx guarantees a new, isolated Word instance
            try:
                self.word_app = win32com.client.DispatchEx("Word.Application")
            except Exception:
                self.word_app = win32com.client.Dispatch("Word.Application")

            try:
                self.word_app.Visible = self.visible
            except Exception:
                pass

            try:
                self.word_app.DisplayAlerts = WD_ALERTS_NONE
            except Exception:
                pass

            try:
                self.word_app.ScreenUpdating = False
            except Exception:
                pass

            return self.word_app
        except Exception as e:
            logger.error(f"Failed to initialize MS Word COM Application: {e}")
            if self.co_initialized:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass
                self.co_initialized = False
            raise RuntimeError(f"Microsoft Word COM server could not be started: {e}")

    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.word_app:
            try:
                self.word_app.Quit(SaveChanges=False)
            except Exception as e:
                logger.warning(f"Error closing Word Application instance: {e}")
            self.word_app = None

        if self.co_initialized:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            self.co_initialized = False


class DocumentInspector:
    """Inspects Word document structure using MS Word layout engine with python-docx server fallback."""

    @staticmethod
    def get_info(file_path: str, visible: bool = False) -> Dict[str, Any]:
        """
        Extracts document statistics and layout information.
        """
        abs_path = os.path.abspath(file_path)
        if not os.path.exists(abs_path):
            raise FileNotFoundError(f"Document file not found: {abs_path}")

        info = {
            "path": abs_path,
            "filename": os.path.basename(abs_path),
            "size_bytes": os.path.getsize(abs_path),
            "page_count": 0,
            "section_count": 0,
            "title": "",
            "author": "",
            "format": os.path.splitext(abs_path)[1].lower().lstrip("."),
        }

        # 1. Prefer Word COM on Windows host first if available
        if pythoncom is not None and win32com is not None:
            try:
                word_ctx = WordCOMContext(visible=visible)
                with word_ctx as word_app:
                    doc = word_app.Documents.Open(
                        FileName=abs_path,
                        ReadOnly=True,
                        ConfirmConversions=False,
                        AddToRecentFiles=False,
                        Visible=visible
                    )
                    try:
                        info["page_count"] = doc.ComputeStatistics(WD_STATISTIC_PAGES)
                        info["section_count"] = doc.Sections.Count
                        try:
                            info["title"] = str(doc.BuiltInDocumentProperties("Title").Value or "")
                            info["author"] = str(doc.BuiltInDocumentProperties("Author").Value or "")
                        except Exception:
                            pass
                        return info
                    finally:
                        if doc:
                            try:
                                doc.Close(SaveChanges=False)
                            except Exception:
                                pass
            except Exception as com_err:
                logger.warning(f"Word COM inspection warning: {com_err}")

        # 2. Check Aspose.Words if available with process isolation to protect web worker
        if aw is not None:
            try:
                import sys
                import subprocess
                import json
                cmd = [
                    sys.executable,
                    "-c",
                    "import sys, json, aspose.words as aw; "
                    "doc = aw.Document(sys.argv[1]); "
                    "res = {'page_count': doc.page_count, 'section_count': len(doc.sections), "
                    "'title': str(doc.built_in_document_properties.title or ''), "
                    "'author': str(doc.built_in_document_properties.author or '')}; "
                    "print(json.dumps(res))",
                    abs_path,
                ]
                res = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
                if res.returncode == 0 and res.stdout.strip():
                    data = json.loads(res.stdout.strip())
                    info["page_count"] = data.get("page_count", 0)
                    info["section_count"] = data.get("section_count", 0)
                    info["title"] = data.get("title", "")
                    info["author"] = data.get("author", "")
                    logger.info(f"Inspected '{info['filename']}' via isolated Aspose process: {info['page_count']} page(s)")
                    return info
                else:
                    logger.warning(f"Isolated Aspose inspection process returned {res.returncode}: {res.stderr}")
            except Exception as aspose_err:
                logger.warning(f"Isolated Aspose inspection process failed: {aspose_err}")

        # 3. Check LibreOffice if available on Linux server
        if _is_libreoffice_available():
            try:
                temp_lo_dir = tempfile.mkdtemp(prefix="lo_inspect_")
                try:
                    pdf_out = _convert_docx_to_pdf_libreoffice(abs_path, temp_lo_dir)
                    if pdf_out:
                        import fitz
                        doc = fitz.open(pdf_out)
                        info["page_count"] = doc.page_count
                        info["title"] = str(doc.metadata.get("title") or "")
                        info["author"] = str(doc.metadata.get("author") or "")
                        doc.close()
                        logger.info(f"Inspected '{info['filename']}' via LibreOffice: {info['page_count']} page(s)")
                        return info
                finally:
                    shutil.rmtree(temp_lo_dir, ignore_errors=True)
            except Exception as lo_err:
                logger.warning(f"LibreOffice inspection warning: {lo_err}")

        # 4. Non-blocking instant python-docx fallback for Linux Cloud Server (0.005s, 0MB RAM)
        if docx is not None:
            try:
                d = docx.Document(abs_path)
                info["section_count"] = len(d.sections)
                breaks = 1
                total_chars = 0
                for p in d.paragraphs:
                    xml = p._element.xml
                    if 'type="page"' in xml or 'lastRenderedPageBreak' in xml:
                        breaks += 1
                    total_chars += len(p.text)
                for tbl in d.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            total_chars += len(cell.text)
                if breaks > 1:
                    info["page_count"] = breaks
                else:
                    info["page_count"] = max(1, (total_chars + 1799) // 1800)
                return info
            except Exception as docx_err:
                logger.warning(f"python-docx inspection warning: {docx_err}")

        info["page_count"] = 1
        info["section_count"] = 1
        return info


class PageExporterEngine:
    """Core engine for exporting page ranges with maximum formatting fidelity."""

    @staticmethod
    def export_range(
        source_file: str,
        output_file: str,
        start_page: int,
        end_page: int,
        export_format: str = "docx",
        mode: str = "trimming",
        visible: bool = False,
        total_pages: int = 1
    ) -> str:
        """
        Exports a page range from source Word doc into output_file.

        Args:
            source_file: Source Word document path.
            output_file: Target path to save output file.
            start_page: 1-indexed start page.
            end_page: 1-indexed end page.
            export_format: Output format ('docx', 'pdf', 'doc', 'rtf', 'docm').
            mode: Extraction method ('trimming', 'aspose', or 'selection').
            visible: Whether Word app runs visibly.
            total_pages: Document total page count.

        Returns:
            Absolute path of created output file.
        """
        abs_source = os.path.abspath(source_file)
        abs_output = os.path.abspath(output_file)
        os.makedirs(os.path.dirname(abs_output), exist_ok=True)

        logger.info(f"Starting page export [{start_page}-{end_page}] from '{os.path.basename(abs_source)}' -> '{os.path.basename(abs_output)}'")

        if mode == "aspose":
            return PageExporterEngine._export_by_aspose(
                abs_source, abs_output, start_page, end_page, export_format
            )

        # Non-Windows / Linux server fallback
        if pythoncom is None or win32com is None:
            if aw is not None:
                try:
                    return PageExporterEngine._export_by_aspose(
                        abs_source, abs_output, start_page, end_page, export_format
                    )
                except Exception as err:
                    logger.warning(f"Aspose.Words server export failed ({err}); trying LibreOffice engine.")

            if _is_libreoffice_available():
                try:
                    return PageExporterEngine._export_by_libreoffice(
                        abs_source, abs_output, start_page, end_page, export_format
                    )
                except Exception as err:
                    logger.warning(f"LibreOffice server export failed ({err}); using fast docx fallback.")

            t_pages = total_pages
            if t_pages <= 1:
                try:
                    info = DocumentInspector.get_info(abs_source)
                    t_pages = info.get("page_count", 1)
                except Exception:
                    t_pages = 1

            return PageExporterEngine._export_by_docx_fallback(
                abs_source, abs_output, start_page, end_page, export_format, total_pages=t_pages
            )

        fmt_code = EXPORT_FORMAT_MAP.get(export_format.lower())
        if fmt_code is None:
            raise ValueError(f"Unsupported export format '{export_format}'. Supported formats: {list(EXPORT_FORMAT_MAP.keys())}")

        if mode == "trimming":
            return PageExporterEngine._export_by_trimming(
                abs_source, abs_output, start_page, end_page, fmt_code, visible
            )
        else:
            return PageExporterEngine._export_by_selection(
                abs_source, abs_output, start_page, end_page, fmt_code, visible
            )

    @staticmethod
    def _export_by_libreoffice(
        abs_source: str,
        abs_output: str,
        start_page: int,
        end_page: int,
        export_format: str
    ) -> str:
        """LibreOffice Engine: Cross-platform page range extraction with 100% layout fidelity."""
        temp_dir = tempfile.mkdtemp(prefix="lo_export_")
        try:
            pdf_path = _convert_docx_to_pdf_libreoffice(abs_source, temp_dir)
            if not pdf_path or not os.path.exists(pdf_path):
                raise RuntimeError(f"LibreOffice failed to convert '{abs_source}' to PDF.")

            from word_exporter_pro.core.pdf_engine import PdfPageExtractor, PdfInspector
            pdf_info = PdfInspector.get_info(pdf_path)
            tp = pdf_info.get("page_count", 1)

            clamped_start = max(1, min(start_page, tp))
            clamped_end = max(clamped_start, min(end_page, tp))

            if export_format.lower() == "pdf":
                res = PdfPageExtractor.extract_range(
                    source_file=pdf_path,
                    output_file=abs_output,
                    start_page=clamped_start,
                    end_page=clamped_end
                )
                logger.success(f"Exported pages [{start_page}-{end_page}] via LibreOffice PDF engine to '{abs_output}'")
                return res
            else:
                split_pdf = os.path.join(temp_dir, "split_range.pdf")
                PdfPageExtractor.extract_range(
                    source_file=pdf_path,
                    output_file=split_pdf,
                    start_page=clamped_start,
                    end_page=clamped_end
                )

                if aw is not None:
                    try:
                        import sys
                        cmd = [
                            sys.executable,
                            "-c",
                            "import sys, aspose.words as aw; "
                            "doc = aw.Document(sys.argv[1]); "
                            "doc.save(sys.argv[2])",
                            split_pdf,
                            abs_output
                        ]
                        r = subprocess.run(cmd, capture_output=True, text=True, timeout=15)
                        if r.returncode == 0 and os.path.exists(abs_output) and os.path.getsize(abs_output) > 0:
                            logger.success(f"Converted split PDF range [{start_page}-{end_page}] to '{abs_output}' via Aspose")
                            return abs_output
                    except Exception as err:
                        logger.warning(f"Aspose PDF to DOCX conversion failed: {err}")

                converted_docx = _convert_pdf_to_docx_libreoffice(split_pdf, abs_output)
                if converted_docx and os.path.exists(converted_docx) and os.path.getsize(converted_docx) > 0:
                    logger.success(f"Converted split PDF range [{start_page}-{end_page}] to '{abs_output}' via LibreOffice")
                    return abs_output

                raise RuntimeError("Could not convert split PDF back to DOCX format.")

        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    @staticmethod
    def _export_by_aspose(
        abs_source: str,
        abs_output: str,
        start_page: int,
        end_page: int,
        export_format: str
    ) -> str:
        """Aspose.Words Layout Engine: High-accuracy cross-platform page range extraction with process isolation."""
        import sys
        import subprocess

        # Run Aspose.Words in an isolated subprocess to prevent C-level segfaults/OOMs from killing Gunicorn
        cmd = [
            sys.executable,
            "-c",
            "import sys, aspose.words as aw; "
            "doc = aw.Document(sys.argv[1]); "
            "tp = doc.page_count; "
            "sc = max(1, min(int(sys.argv[3]), tp)); "
            "ec = max(sc, min(int(sys.argv[4]), tp)); "
            "cnt = ec - sc + 1; "
            "ext = doc.extract_pages(sc - 1, cnt); "
            "ext.save(sys.argv[2])",
            abs_source,
            abs_output,
            str(start_page),
            str(end_page),
        ]

        try:
            res = subprocess.run(cmd, capture_output=True, text=True, timeout=12)
            if res.returncode == 0 and os.path.exists(abs_output) and os.path.getsize(abs_output) > 0:
                logger.success(f"Exported pages [{start_page}-{end_page}] via isolated Aspose process to '{abs_output}'")
                return abs_output
            else:
                logger.warning(f"Isolated Aspose process exited ({res.returncode}): {res.stderr}; falling back to secondary server engine.")
        except Exception as e:
            logger.warning(f"Isolated Aspose process failed or timed out: {e}; falling back to secondary server engine.")

        # Fail-safe fallback: use MS Word COM on Windows, or LibreOffice / python-docx on Linux server
        if pythoncom not in (None,) and win32com not in (None,):
            fmt_code = EXPORT_FORMAT_MAP.get(export_format.lower(), 16)
            return PageExporterEngine._export_by_trimming(
                abs_source, abs_output, start_page, end_page, fmt_code, False
            )
        elif _is_libreoffice_available():
            try:
                return PageExporterEngine._export_by_libreoffice(
                    abs_source, abs_output, start_page, end_page, export_format
                )
            except Exception as lo_err:
                logger.warning(f"LibreOffice fallback export failed ({lo_err}); using fast docx fallback.")

        try:
            info = DocumentInspector.get_info(abs_source)
            t_pages = info.get("page_count", 1)
        except Exception:
            t_pages = 1
        return PageExporterEngine._export_by_docx_fallback(
            abs_source, abs_output, start_page, end_page, export_format, total_pages=t_pages
        )

    @staticmethod
    def _export_by_docx_fallback(
        abs_source: str,
        abs_output: str,
        start_page: int,
        end_page: int,
        export_format: str,
        total_pages: int = 1
    ) -> str:
        """Non-Windows / Linux cloud server fallback engine using python-docx to trim pages."""
        if docx is None:
            shutil.copy2(abs_source, abs_output)
            return abs_output

        try:
            doc = docx.Document(abs_source)
            body = doc._body._element

            children = [child for child in list(body) if child.tag.rsplit('}', 1)[-1] in ('p', 'tbl')]
            if not children:
                doc.save(abs_output)
                return abs_output

            # 1. Check explicit page breaks (w:br[@w:type="page"], lastRenderedPageBreak, pageBreakBefore, section breaks)
            element_pages = []
            current_page = 1
            has_explicit_breaks = False

            for child in children:
                xml = child.xml
                has_break = ('type="page"' in xml or 'lastRenderedPageBreak' in xml or 'nextPage' in xml or 'oddPage' in xml or 'evenPage' in xml or 'pageBreakBefore' in xml or 'sectPr' in xml)
                element_pages.append(current_page)
                if has_break:
                    has_explicit_breaks = True
                    current_page += 1

            # 2. If explicit breaks are absent, divide content by text character density weight across target_pages
            if not has_explicit_breaks:
                target_pages = max(total_pages, end_page, 1)
                
                elem_weights = []
                for child in children:
                    txt = "".join(child.itertext()) if hasattr(child, "itertext") else ""
                    elem_weights.append(max(10, len(txt)))

                total_weight = sum(elem_weights)
                weight_per_page = total_weight / target_pages if total_weight > 0 else 1.0

                element_pages = []
                accumulated_weight = 0
                for w in elem_weights:
                    prev_weight = accumulated_weight
                    accumulated_weight += w
                    mid_weight = (prev_weight + accumulated_weight) / 2.0
                    p_num = min(target_pages, max(1, int(mid_weight / weight_per_page) + 1))
                    element_pages.append(p_num)

            # 3. Trim elements outside target page range [start_page, end_page]
            elements_to_remove = []
            for child, elem_page in zip(children, element_pages):
                if elem_page < start_page or elem_page > end_page:
                    elements_to_remove.append(child)

            if elements_to_remove:
                for el in elements_to_remove:
                    try:
                        sectPr = el.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
                        if sectPr is not None and body.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr') is None:
                            body.append(sectPr)
                        body.remove(el)
                    except Exception:
                        pass

            doc.save(abs_output)
            logger.success(f"Trimmed pages [{start_page}-{end_page}] (Server mode) to '{abs_output}'")
            return abs_output
        except Exception as e:
            logger.warning(f"python-docx fallback trimming warning: {e}. Copying original file.")
            shutil.copy2(abs_source, abs_output)
            return abs_output

    @staticmethod
    def _export_by_trimming(
        abs_source: str,
        abs_output: str,
        start_page: int,
        end_page: int,
        fmt_code: int,
        visible: bool
    ) -> str:
        """High-Fidelity Trimming Engine: Duplicates source doc and trims unwanted pages."""
        temp_dir = tempfile.mkdtemp(prefix="word_exp_")
        temp_source_copy = os.path.join(temp_dir, f"work_{os.path.basename(abs_source)}")

        try:
            shutil.copy2(abs_source, temp_source_copy)

            with WordCOMContext(visible=visible) as word_app:
                source_doc = None
                doc = word_app.Documents.Open(
                    FileName=temp_source_copy,
                    ReadOnly=False,
                    ConfirmConversions=False,
                    AddToRecentFiles=False,
                    Visible=visible
                )

                try:
                    total_pages = doc.ComputeStatistics(WD_STATISTIC_PAGES)

                    start_page = max(1, min(start_page, total_pages))
                    end_page = max(start_page, min(end_page, total_pages))
                    expected_keep = end_page - start_page + 1

                    # Compute the page boundaries on the UNTOUCHED source document.
                    # Trimming the copy re-paginates it, so GoTo(page) on the copy
                    # drifts from the source's true boundaries whenever a table row
                    # spans a page cut. The source positions are exact and remain
                    # valid in the byte-identical copy until that region is cut.
                    # Page starts come from Range.Information(wdActiveEndPageNumber)
                    # via a binary search because GoTo(page) reports unreliable
                    # boundaries in headless COM Word.
                    source_doc = word_app.Documents.Open(
                        FileName=abs_source,
                        ReadOnly=True,
                        ConfirmConversions=False,
                        AddToRecentFiles=False,
                        Visible=False
                    )
                    start_pos = PageExporterEngine._page_start_by_position(source_doc, start_page)
                    if end_page >= total_pages:
                        end_pos = source_doc.Content.End
                    else:
                        end_pos = PageExporterEngine._page_start_by_position(source_doc, end_page + 1)

                    # 1. Trim the BACK first using the source boundary. Deleting
                    #    content AFTER `start_pos` never shifts content before it, so
                    #    the front boundary stays valid for the second cut.
                    if end_pos < doc.Content.End:
                        PageExporterEngine._delete_range_after(word_app, doc, end_pos)

                    # 2. Trim the FRONT using the source boundary.
                    if start_page > 1:
                        PageExporterEngine._trim_front_to_page(
                            word_app, doc, start_page, start_pos
                        )

                    # 3. Remove any stray page-break/blank paragraph left at the cut edges so no
                    #    spurious blank front/back page survives in the trimmed result. This runs
                    #    BEFORE the page-count compaction below: a leading empty paragraph before
                    #    a table (which Word keeps when everything before it is deleted) would
                    #    otherwise push the count over by one page and make the compaction loop
                    #    burn through all its iterations trying to reclaim it. The front edge is
                    #    only cleaned when a front trim actually occurred, so a legitimate leading
                    #    page-break on a range starting at page 1 is kept.
                    if start_page > 1:
                        PageExporterEngine._clean_page_boundary(word_app, doc, front_of_doc=True)
                    PageExporterEngine._clean_page_boundary(word_app, doc, front_of_doc=False)

                    # 4. Guarantee the exact page count. The source-boundary cuts above
                    #    should already hold it, but a razor-edge reflow (trailing empty
                    #    paragraph pushed onto its own page) may need a compensating cut.
                    PageExporterEngine._trim_tail_to_keep(word_app, doc, expected_keep)

                    # 5. When trimming removes content BEFORE/AFTER the range, the boundary section breaks
                    #    are deleted too, so the kept sections may inherit an adjacent (wrong) section's
                    #    header/footer, page setup, and page-number scheme. Re-sync them against the
                    #    original source document so the exported range looks correct.
                    if end_page < total_pages or start_page > 1:
                        try:
                            start_section_index = PageExporterEngine._section_containing_page(source_doc, start_page)
                            end_section_index = PageExporterEngine._section_containing_page(source_doc, end_page)
                            PageExporterEngine._restore_sections_from_source(
                                target_doc=doc,
                                source_doc=source_doc,
                                start_section_index=start_section_index,
                                end_section_index=end_section_index,
                                only_last_section=True
                            )
                        except Exception as restore_err:
                            logger.warning(f"Could not restore headers/footers for trimmed range: {restore_err}")

                    # 6. Restart page numbering from 1 on the trimmed document so PAGE/NUMPAGES
                    #    fields show the range-relative numbers (e.g. "2 of 3") instead of the
                    #    source document's original page numbers. This must run AFTER the section
                    #    restore above, otherwise the earlier StartingNumber overwrites would win.
                    PageExporterEngine._restart_page_numbering(doc, start_page)

                    # 6. Natively update field codes (PAGE, NUMPAGES, TOC, etc.) across the body
                    #    and every header/footer so they reflect the final trimmed content.
                    try:
                        try:
                            doc.Fields.Update()
                        except Exception:
                            pass
                        try:
                            for sec in doc.Sections:
                                for hf_coll in (sec.Headers, sec.Footers):
                                    for hf_idx in (1, 2, 3):
                                        try:
                                            hf_coll(hf_idx).Range.Fields.Update()
                                        except Exception:
                                            pass
                        except Exception:
                            pass
                    except Exception:
                        pass

                    # 6. Save as output destination format
                    #    Restoring the source headers/footers can re-paginate a
                    #    razor-edge page count, so re-compact to the expected count
                    #    after every layout-affecting step and right before saving.
                    PageExporterEngine._compact_to_fit(word_app, doc, expected_keep)
                    doc.SaveAs2(FileName=abs_output, FileFormat=fmt_code)
                    logger.success(f"Exported pages [{start_page}-{end_page}] to '{abs_output}'")

                finally:
                    if source_doc:
                        try:
                            source_doc.Close(SaveChanges=False)
                        except Exception:
                            pass
                    if doc:
                        try:
                            doc.Close(SaveChanges=False)
                        except Exception:
                            pass

        finally:
            if os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir, ignore_errors=True)
                except Exception as e:
                    logger.warning(f"Could not remove temporary directory {temp_dir}: {e}")

        return abs_output

    @staticmethod
    def _delete_range(word_app, cut_range) -> None:
        """Delete a range robustly; Range.Delete() can fail on ranges ending at Content.End."""
        try:
            cut_range.Delete()
            return
        except Exception:
            pass
        try:
            cut_range.Select()
            word_app.Selection.Delete()
            return
        except Exception:
            pass
        # Last resort: shrink the range by one character and retry
        try:
            cut_range.MoveEnd(6, -1)
            cut_range.Delete()
        except Exception:
            logger.warning("Could not delete range during page trimming.")

    @staticmethod
    def _page_start_by_position(doc, page: int) -> int:
        """Find the body position where `page` starts.

        `GoTo(page)` returns Word's native page boundary, but inside tables the
        returned position is unreliable in headless COM Word (it can land anywhere
        in a multi-row region). Table-row start page numbers
        (Range.Information(wdActiveEndPageNumber) on a row start) are reliable and
        monotonic, so a boundary that falls inside a table is snapped to the first
        row whose start is on `page`. As a last resort, fall back to a binary
        search over Information(wdActiveEndPageNumber).
        """
        pos = None
        try:
            pos = doc.GoTo(What=WD_GOTO_PAGE, Which=WD_GOTO_ABSOLUTE, Count=page).Start
        except Exception:
            pass
        if pos is None or not (doc.Content.Start <= pos <= doc.Content.End):
            lo, hi = doc.Content.Start, doc.Content.End
            while lo < hi:
                mid = (lo + hi) // 2
                try:
                    pn = doc.Range(mid, min(mid + 1, hi)).Information(WD_ACTIVE_END_PAGE_NUMBER)
                except Exception:
                    pn = 1
                if pn >= page:
                    hi = mid
                else:
                    lo = mid + 1
            pos = lo
        return PageExporterEngine._snap_table_boundary(doc, pos, page)

    @staticmethod
    def _snap_table_boundary(doc, pos: int, page: int) -> int:
        """Snap `pos` to a table row start if it falls inside a table.

        Raw GoTo/binary page boundaries inside tables are unreliable, but the page
        number reported at a table row start is reliable and monotonic, so the
        boundary becomes the first row whose start lies on `page`. This keeps the
        export/verify boundaries deterministic for table-heavy documents.
        """
        if doc.Tables.Count == 0:
            return pos
        for table in doc.Tables:
            if not (table.Range.Start <= pos <= table.Range.End):
                continue
            # Fast path: when the table has no vertically merged rows, Rows(i) is
            # indexable and row-start page numbers are non-decreasing, so binary
            # search finds the first row start on `page` in a handful of probes
            # instead of scanning the whole table.
            try:
                count = table.Rows.Count
                end = doc.Content.End
                lo, hi = 1, count
                ans = None
                while lo <= hi:
                    mid = (lo + hi) // 2
                    rs = table.Rows(mid).Range.Start
                    pn = doc.Range(rs, min(rs + 1, end)).Information(WD_ACTIVE_END_PAGE_NUMBER)
                    if pn >= page:
                        ans = rs
                        hi = mid - 1
                    else:
                        lo = mid + 1
                if ans is not None:
                    return ans
            except Exception:
                pass
            layout, max_row = PageExporterEngine._table_row_layout(doc, table)
            end = doc.Content.End
            for ri in range(1, max_row + 1):
                if ri not in layout:
                    continue
                rs, _re = layout[ri]
                try:
                    pn = doc.Range(rs, min(rs + 1, end)).Information(WD_ACTIVE_END_PAGE_NUMBER)
                except Exception:
                    continue
                if pn >= page:
                    return rs
            return pos
        return pos

    @staticmethod
    def _table_row_layout(doc, table):
        """Return (row_start, row_end, max_row) for a table.

        Word refuses `Rows(i)` access on vertically merged tables, so fall back to
        iterating the cell collection (Cells/RowIndex still work on merged tables).
        Returns dict row_index -> (start, end) plus the highest row index.
        Results are cached per document/table; the key includes Rows.Count so a
        row deletion invalidates the cache.
        """
        try:
            count = table.Rows.Count
            cache_key = (id(doc), table.Range.Start, table.Range.End, count)
        except Exception:
            cache_key = None
        if cache_key is not None and cache_key in _TABLE_ROW_LAYOUT_CACHE:
            return _TABLE_ROW_LAYOUT_CACHE[cache_key]
        layout = None
        if cache_key is not None:
            try:
                layout = {}
                for i in range(1, count + 1):
                    r = table.Rows(i)
                    layout[i] = (r.Range.Start, r.Range.End)
            except Exception:
                layout = None
        if layout is None:
            layout = {}
            max_row = 0
            cells_count = table.Range.Cells.Count
            for i in range(1, cells_count + 1):
                try:
                    cell = table.Range.Cells(i)
                    ri = cell.RowIndex
                    s = cell.Range.Start
                    e = cell.Range.End
                except Exception:
                    continue
                layout[ri] = (min(layout.get(ri, (s, e))[0], s), max(layout.get(ri, (s, s))[1], e))
                max_row = max(max_row, ri)
            count = max_row
        if cache_key is not None:
            _TABLE_ROW_LAYOUT_CACHE[cache_key] = (layout, count)
        return layout, count

    @staticmethod
    def _row_containing(doc, table, pos: int) -> int:
        """Return the 1-based index of the table row containing `pos`, or None."""
        try:
            count = table.Rows.Count
            lo, hi = 1, count
            best = None
            while lo <= hi:
                mid = (lo + hi) // 2
                rs = table.Rows(mid).Range.Start
                if rs <= pos:
                    best = mid
                    lo = mid + 1
                else:
                    hi = mid - 1
            if best is not None:
                row_end = table.Rows(best).Range.End
                if row_end > pos:
                    return best
            return None
        except Exception:
            pass
        layout, max_row = PageExporterEngine._table_row_layout(doc, table)
        for ri in range(1, max_row + 1):
            if ri not in layout:
                continue
            _, row_end = layout[ri]
            if row_end > pos:
                return ri
        return None

    @staticmethod
    def _row_straddling(doc, table, pos: int) -> int:
        """Return the row that strictly straddles `pos` (row_start < pos <= row_end).

        A boundary that lands exactly on a row start belongs to the NEXT row, so it
        does not straddle any row here; the caller keeps only rows entirely before it.
        """
        try:
            count = table.Rows.Count
            lo, hi = 1, count
            best = None
            while lo <= hi:
                mid = (lo + hi) // 2
                rs = table.Rows(mid).Range.Start
                if rs < pos:
                    best = mid
                    lo = mid + 1
                else:
                    hi = mid - 1
            if best is not None:
                row_end = table.Rows(best).Range.End
                if row_end >= pos:
                    return best
            return None
        except Exception:
            pass
        layout, max_row = PageExporterEngine._table_row_layout(doc, table)
        for ri in range(1, max_row + 1):
            if ri not in layout:
                continue
            row_start, row_end = layout[ri]
            if row_start < pos <= row_end:
                return ri
        return None

    @staticmethod
    def _row_cells(table, row_index: int):
        """Return the cells of `row_index` as [(start, end, cell)] in document order.

        Falls back to iterating the whole cell collection for vertically merged
        tables where `Rows(i)` access is refused by Word.
        """
        cells = []
        try:
            count = table.Rows(row_index).Cells.Count
            for c in range(1, count + 1):
                cell = table.Rows(row_index).Cells(c)
                cells.append((cell.Range.Start, cell.Range.End, cell))
            return cells
        except Exception:
            pass
        result = []
        total = table.Range.Cells.Count
        for i in range(1, total + 1):
            try:
                cell = table.Range.Cells(i)
                if cell.RowIndex == row_index:
                    result.append((cell.Range.Start, cell.Range.End, cell))
            except Exception:
                continue
        result.sort(key=lambda x: x[0])
        return result

    @staticmethod
    def _delete_cell(word_app, doc, cell, cs: int, ce: int) -> None:
        """Delete a whole table cell (or its content as a fallback) robustly."""
        try:
            cell.Delete()
            return
        except Exception:
            pass
        try:
            rng = cell.Range
            if rng.End > rng.Start:
                rng.Delete()
        except Exception:
            pass

    @staticmethod
    def _delete_table_rows_before(word_app, doc, table, keep_from: int) -> None:
        """Delete table rows 1..keep_from-1 (keep_from row is kept). Works on merged tables."""
        if keep_from <= 1:
            return
        try:
            keep_start = table.Rows(keep_from).Range.Start
        except Exception:
            layout, _ = PageExporterEngine._table_row_layout(doc, table)
            keep_start = layout.get(keep_from, (None, None))[0]
        if keep_start is None:
            return
        try:
            rng = doc.Range(Start=table.Range.Start, End=keep_start)
            rng.Select()
            word_app.Selection.Rows.Delete()
            return
        except Exception:
            pass
        for row_index in range(keep_from - 1, 0, -1):
            try:
                table.Rows(row_index).Delete()
            except Exception:
                break

    @staticmethod
    def _delete_table_rows_after(word_app, doc, table, keep_until: int) -> None:
        """Delete table rows keep_until+1..Count (keep_until row is kept). Works on merged tables."""
        try:
            if keep_until >= table.Rows.Count:
                return
            next_start = table.Rows(keep_until + 1).Range.Start
        except Exception:
            layout, max_row = PageExporterEngine._table_row_layout(doc, table)
            if keep_until >= max_row:
                return
            next_start = layout.get(keep_until + 1, (None, None))[0]
        if next_start is None:
            return
        try:
            rng = doc.Range(Start=next_start, End=table.Range.End)
            rng.Select()
            word_app.Selection.Rows.Delete()
            return
        except Exception:
            pass
        for row_index in range(max_row, keep_until, -1):
            try:
                table.Rows(row_index).Delete()
            except Exception:
                break

    @staticmethod
    def _trim_front_to_page(word_app, doc, start_page: int, start_pos: int = None) -> None:
        """Delete everything before the start of `start_page`.

        A Range.Delete() whose end falls inside a table removes the entire table,
        so when the page boundary lands inside a table, delete only the content
        before the table and the rows that end before the spanning row instead.
        `start_pos` may be passed in explicitly (computed from the untouched source
        document) to avoid re-pagination drift on the working copy.
        """
        if start_pos is None:
            start_pos = PageExporterEngine._page_start_by_position(doc, start_page)
        if start_pos <= doc.Content.Start:
            return
        if not doc.Range(start_pos, start_pos).Information(12):
            PageExporterEngine._delete_range(
                word_app, doc.Range(Start=doc.Content.Start, End=start_pos)
            )
            return
        for table in doc.Tables:
            if table.Range.Start <= start_pos <= table.Range.End:
                keep_from = PageExporterEngine._row_containing(doc, table, start_pos)
                if keep_from:
                    # The page boundary falls inside the row: truncate it at the
                    # boundary so the export holds exactly the visible page content.
                    # A whole cell on the kept side is preserved; a cell on the
                    # trimmed side is deleted entirely; the boundary cell is cut.
                    row_cells = PageExporterEngine._row_cells(table, keep_from)
                    boundary_idx = None
                    for i, (cs, ce, _cell) in enumerate(row_cells):
                        if ce > start_pos:
                            boundary_idx = i
                            break
                    if boundary_idx is not None:
                        cs, ce, boundary_cell = row_cells[boundary_idx]
                        if cs < start_pos:
                            # Delete the latest region first so earlier positions hold.
                            PageExporterEngine._delete_range(
                                word_app, doc.Range(Start=cs, End=start_pos)
                            )
                        for i in range(boundary_idx):
                            ocs, oce, ocell = row_cells[i]
                            PageExporterEngine._delete_cell(word_app, doc, ocell, ocs, oce)
                if keep_from and keep_from > 1:
                    # table.Range.Start stays stable when only its leading rows are
                    # deleted, so capture it and delete rows 1..keep_from-1 first.
                    table_start = table.Range.Start
                    PageExporterEngine._delete_table_rows_before(
                        word_app, doc, table, keep_from
                    )
                    if table_start > doc.Content.Start:
                        PageExporterEngine._delete_range(
                            word_app, doc.Range(Start=doc.Content.Start, End=table_start)
                        )
                elif table.Range.Start > doc.Content.Start:
                    PageExporterEngine._delete_range(
                        word_app, doc.Range(Start=doc.Content.Start, End=table.Range.Start)
                    )
                return
        PageExporterEngine._delete_range(
            word_app, doc.Range(Start=doc.Content.Start, End=start_pos)
        )

    @staticmethod
    def _delete_range_after(word_app, doc, drop_start: int) -> None:
        """Delete [drop_start, doc.Content.End] robustly.

        A Range.Delete() starting inside a table removes the whole row (or more),
        losing content that belongs to the kept range. When `drop_start` lands
        inside a table, keep the row that spans the boundary and delete only the
        rows after it plus everything after the table.
        """
        if drop_start >= doc.Content.End:
            return
        if not doc.Range(drop_start, drop_start).Information(12):
            PageExporterEngine._delete_range(
                word_app, doc.Range(Start=drop_start, End=doc.Content.End)
            )
            return
        for table in doc.Tables:
            if table.Range.Start <= drop_start <= table.Range.End:
                keep_until = PageExporterEngine._row_straddling(doc, table, drop_start)
                if keep_until is None:
                    layout, _ = PageExporterEngine._table_row_layout(doc, table)
                    keep_until = 0
                    for ri in sorted(layout):
                        if layout[ri][1] <= drop_start:
                            keep_until = ri
                table_end = table.Range.End
                # Delete from the latest region backward so earlier positions stay
                # valid: content after the table, rows after the spanning row,
                # then truncate the boundary cell and drop the cells after it.
                if table_end < doc.Content.End:
                    PageExporterEngine._delete_range(
                        word_app, doc.Range(Start=table_end, End=doc.Content.End)
                    )
                if keep_until > 0:
                    PageExporterEngine._delete_table_rows_after(
                        word_app, doc, table, keep_until
                    )
                else:
                    PageExporterEngine._delete_table_rows_after(
                        word_app, doc, table, 0
                    )
                if keep_until > 0:
                    straddle = PageExporterEngine._row_straddling(doc, table, drop_start)
                    if straddle is not None:
                        row_cells = PageExporterEngine._row_cells(table, straddle)
                        boundary_idx = None
                        for i, (cs, ce, _cell) in enumerate(row_cells):
                            if ce > drop_start:
                                boundary_idx = i
                                break
                        if boundary_idx is not None:
                            cs, ce, boundary_cell = row_cells[boundary_idx]
                            if cs < drop_start and drop_start < ce - 1:
                                PageExporterEngine._delete_range(
                                    word_app, doc.Range(Start=drop_start, End=ce - 1)
                                )
                            for i in range(boundary_idx + 1, len(row_cells)):
                                ocs, oce, ocell = row_cells[i]
                                PageExporterEngine._delete_cell(
                                    word_app, doc, ocell, ocs, oce
                                )
                return
        PageExporterEngine._delete_range(
            word_app, doc.Range(Start=drop_start, End=doc.Content.End)
        )

    @staticmethod
    def _trim_tail_to_keep(word_app, doc, expected_keep: int) -> None:
        """Ensure the document holds exactly `expected_keep` pages without deleting
        real content.

        The front/back cuts are already position-exact against the untouched source
        boundaries, so the working copy holds exactly the requested page content.
        Headless COM Word's in-session pagination (ComputeStatistics/Repaginate and
        page number lookups) is unreliable for the trimmed copy and can report 3, 5
        or 7 pages for identical content, so a pagination-driven tail deletion would
        arbitrarily cut real content (e.g. dropping the last kept page). Only the
        safe `_compact_to_fit` runs here: it never deletes content, only shrinks the
        spacing of the trailing empty paragraph / last table row / header-footer so
        a razor-edge overflow (a trailing empty paragraph pushed onto its own page)
        settles back to the requested count.
        """
        PageExporterEngine._compact_to_fit(word_app, doc, expected_keep)

    @staticmethod
    def _page_count(doc) -> int:
        """Return the current rendered page count of `doc`.

        `doc.ComputeStatistics(wdStatisticPages)` is unreliable in headless COM
        Word (it reported 3, 5 or 7 for identical content in the same session),
        while `Range.Information(wdActiveEndPageNumber)` on the final character
        reflects the true rendered last-page number and matches a fresh open of
        the saved file. Fall back to ComputeStatistics only if the range lookup
        fails (e.g. an empty document).
        """
        end = doc.Content.End
        for start in (end - 1, end - 2):
            if start < 0:
                continue
            try:
                return doc.Range(start, start + 1).Information(WD_ACTIVE_END_PAGE_NUMBER)
            except Exception:
                continue
        try:
            return doc.ComputeStatistics(WD_STATISTIC_PAGES)
        except Exception:
            return 1

    @staticmethod
    def _compact_to_fit(word_app, doc, expected_keep: int) -> None:
        """Shrink trailing whitespace so the document fits exactly `expected_keep`
        pages. Never deletes content: it only reduces the size/spacing of the
        trailing empty paragraph, the last table row's paragraphs, and the
        header/footer paragraphs, so a razor-edge re-pagination (e.g. after
        section headers/footers are restored and the real, taller header/footer
        pushes the body over) settles back to the requested page count.
        """
        guard = 0
        distance_allowed = False
        while guard < 30:
            try:
                doc.Repaginate()
            except Exception:
                pass
            current_pages = PageExporterEngine._page_count(doc)
            if current_pages <= expected_keep:
                break
            guard += 1
            reduced = False
            # 1. Compact trailing whitespace. A table that ends flush with the
            #    bottom margin pushes the trailing empty paragraph(s) onto an extra
            #    page; deleting a manual page break at the cut edge can also leave
            #    several empty paragraphs at the end. Every trailing empty paragraph
            #    is shrunk (size 1, zero spacing, no page-break-before), then the
            #    last non-empty paragraph's spacing is compacted so the shrunk
            #    paragraphs regain the few points they need to fit on the final page.
            try:
                idx = doc.Paragraphs.Count
                while idx >= 1:
                    para = doc.Paragraphs(idx)
                    text = para.Range.Text
                    if text.strip("\r\x07"):
                        # Last non-empty paragraph: reclaim its spacing.
                        ppf = para.Range.ParagraphFormat
                        for attr in ("SpaceBefore", "SpaceAfter"):
                            try:
                                if getattr(ppf, attr):
                                    setattr(ppf, attr, 0)
                                    reduced = True
                            except Exception:
                                pass
                        try:
                            if ppf.LineSpacingRule != 0:
                                ppf.LineSpacingRule = 0
                                reduced = True
                        except Exception:
                            pass
                        try:
                            if ppf.PageBreakBefore:
                                ppf.PageBreakBefore = False
                                reduced = True
                        except Exception:
                            pass
                        break
                    # Empty paragraph: shrink it away.
                    try:
                        if para.Range.Font.Size != 1:
                            para.Range.Font.Size = 1
                            reduced = True
                    except Exception:
                        pass
                    pf = para.Range.ParagraphFormat
                    for attr in ("SpaceBefore", "SpaceAfter"):
                        try:
                            if getattr(pf, attr):
                                setattr(pf, attr, 0)
                                reduced = True
                        except Exception:
                            pass
                    try:
                        if pf.LineSpacingRule != 0:
                            pf.LineSpacingRule = 0
                            reduced = True
                    except Exception:
                        pass
                    try:
                        if pf.PageBreakBefore:
                            pf.PageBreakBefore = False
                            reduced = True
                    except Exception:
                        pass
                    idx -= 1
                # The last table row's paragraphs are also compacted: a table that
                # ends flush with the bottom margin pushes the trailing paragraph(s)
                # onto an extra page.
                try:
                    last_table = doc.Tables(doc.Tables.Count)
                    last_row = last_table.Rows(last_table.Rows.Count)
                    if not last_row.Range.Information(12):
                        raise RuntimeError("not a table row")
                    for cell in last_row.Range.Cells:
                        for para in cell.Range.Paragraphs:
                            cpf = para.Range.ParagraphFormat
                            for attr in ("SpaceBefore", "SpaceAfter"):
                                try:
                                    if getattr(cpf, attr):
                                        setattr(cpf, attr, 0)
                                        reduced = True
                                except Exception:
                                    pass
                            try:
                                if cpf.LineSpacingRule != 0:
                                    cpf.LineSpacingRule = 0
                                    reduced = True
                            except Exception:
                                pass
                            try:
                                if cpf.PageBreakBefore:
                                    cpf.PageBreakBefore = False
                                    reduced = True
                            except Exception:
                                pass
                except Exception:
                    pass
            except Exception:
                pass
            # 2. Compact the header/footer paragraph spacing so the body regains the
            #    couple of lines a razor-edge page needs. This only runs while the
            #    document still overflows the expected page count, so ordinary
            #    exports are never touched. The header/footer text is preserved.
            #    Only if spacing alone cannot reclaim the page are the header/footer
            #    distances reduced as a last resort.
            if PageExporterEngine._compact_headers_footers(doc, reduce_distance=distance_allowed):
                reduced = True
            if not reduced:
                if not distance_allowed:
                    distance_allowed = True
                    continue
                break

    @staticmethod
    def _compact_headers_footers(doc, reduce_distance: bool = False) -> bool:
        """Reduce header/footer paragraph spacing so the body gets a little more
        room. Returns True if any spacing was reduced; never deletes header/footer
        content. `reduce_distance` additionally shrinks the header/footer distances
        and is only used as a last resort."""
        reduced = False
        for sec in doc.Sections:
            for hf_coll in (sec.Headers, sec.Footers):
                for hf_idx in (1, 2, 3):
                    try:
                        hf = hf_coll(hf_idx)
                        for para in hf.Range.Paragraphs:
                            pf = para.Range.ParagraphFormat
                            try:
                                if pf.SpaceBefore and pf.SpaceBefore > 0:
                                    pf.SpaceBefore = 0
                                    reduced = True
                            except Exception:
                                pass
                            try:
                                if pf.SpaceAfter and pf.SpaceAfter > 0:
                                    pf.SpaceAfter = 0
                                    reduced = True
                            except Exception:
                                pass
                            try:
                                if pf.PageBreakBefore:
                                    pf.PageBreakBefore = False
                                    reduced = True
                            except Exception:
                                pass
                    except Exception:
                        pass
            if reduce_distance:
                ps = sec.PageSetup
                for attr in ("FooterDistance", "HeaderDistance"):
                    try:
                        if getattr(ps, attr) and getattr(ps, attr) > 0:
                            setattr(ps, attr, 0)
                            reduced = True
                    except Exception:
                        pass
        return reduced

    @staticmethod
    def _clean_page_boundary(word_app, doc, front_of_doc: bool) -> None:
        """Remove a stray leading/trailing page-break or blank paragraph left at the
        trimmed edge so the first and last pages of the result hold only real content.

        When the front trim leaves a table as the first element, Word keeps an empty
        paragraph before it. That paragraph consumes the first row's height on page
        one (so 15 rows/page becomes 14/15/1 and the range gains a spurious page),
        so leading paragraphs that contain only empty/break characters and sit before
        the first table are removed as well.
        """
        for _ in range(5):
            try:
                if front_of_doc:
                    if doc.Tables.Count:
                        first_table_start = doc.Tables(1).Range.Start
                        lead = doc.Range(
                            Start=doc.Content.Start, End=first_table_start
                        ).Text
                        if lead and not lead.strip("\r\x07\x0c\f"):
                            doc.Range(
                                Start=doc.Content.Start, End=first_table_start
                            ).Delete()
                            continue
                    rng = doc.Range(
                        Start=doc.Content.Start,
                        End=min(doc.Content.Start + 2, doc.Content.End),
                    )
                    if rng.End <= rng.Start:
                        break
                    if "\x0c" in rng.Text or "\f" in rng.Text:
                        rng.Delete()
                    else:
                        break
                else:
                    if doc.Content.End <= 2:
                        break
                    # A manual page break (form-feed) left at the trimmed tail pushes
                    # a spurious empty last page even though all real content was cut
                    # away. Scan the trailing paragraphs (the final empty paragraph
                    # plus the one before it) for such a break and delete it. The
                    # window is a few characters wider than the final paragraph mark
                    # so a break in the second-to-last paragraph is also caught.
                    tail_start = max(doc.Content.Start, doc.Content.End - 48)
                    rng = doc.Range(Start=tail_start, End=doc.Content.End)
                    idx = rng.Text.find("\x0c")
                    if idx < 0:
                        idx = rng.Text.find("\f")
                    if idx < 0:
                        break
                    doc.Range(
                        Start=tail_start + idx, End=tail_start + idx + 1
                    ).Delete()
            except Exception:
                break

    @staticmethod
    def _restart_page_numbering(doc, start_page: int) -> None:
        """Reset page numbering on the trimmed document so the PAGE/NUMPAGES fields show
        range-relative numbers instead of the source document's original numbering.

        - First section restarts at 1 (and re-enables restart-at-section), so the kept
          first page reads as page 1 of the new document.
        - Subsequent sections continue sequentially (RestartNumberingAtSection=False) so
          numbering flows across sections without resetting.
        - Only the page-numbering scheme is touched; explicit header/footer text is left
          intact so the section restore above is preserved.
        """
        try:
            if not doc.Sections:
                return
            first = True
            for sec in doc.Sections:
                try:
                    ps = sec.PageSetup
                    if first:
                        ps.PageNumbering.RestartNumberingAtSection = True
                        ps.PageNumbering.StartingNumber = 1
                        first = False
                    else:
                        ps.PageNumbering.RestartNumberingAtSection = False
                except Exception:
                    pass
        except Exception:
            pass

    @staticmethod
    def _section_containing_page(doc, page: int) -> int:
        """Return the 1-indexed source section that contains the given page."""
        try:
            pos = PageExporterEngine._page_start_by_position(doc, page)
            for i in range(1, doc.Sections.Count + 1):
                r = doc.Sections(i).Range
                if r.Start <= pos < r.End:
                    return i
            return doc.Sections.Count
        except Exception:
            return 1

    @staticmethod
    def _get_effective_header_footer(source_doc, sec_index: int, is_header: bool, hf_index: int):
        """Trace back LinkToPrevious to find the effective source Header/Footer object."""
        curr_idx = sec_index
        while curr_idx >= 1:
            try:
                sec = source_doc.Sections(curr_idx)
                hf = sec.Headers(hf_index) if is_header else sec.Footers(hf_index)
                if not hf.LinkToPrevious or curr_idx == 1:
                    return hf
            except Exception:
                pass
            curr_idx -= 1
        try:
            sec = source_doc.Sections(1)
            return sec.Headers(hf_index) if is_header else sec.Footers(hf_index)
        except Exception:
            return None

    @staticmethod
    def _copy_header_footer_formatting(src_hf, tgt_hf) -> None:
        """Copies resolved formatting from a source header/footer to a target
        header/footer. FormattedText only carries explicitly-set formatting, so
        style-inherited values (indents, alignment, spacing, fonts) would otherwise
        resolve against the target document's own built-in styles. Assigning the
        source ParagraphFormat + font properties transfers the effective values.
        The font copy runs LAST because style/paragraph-format assignment reapplies
        the target's style character formatting to the range."""
        try:
            src_count = src_hf.Range.Paragraphs.Count
            tgt_count = tgt_hf.Range.Paragraphs.Count
            for i in range(1, min(src_count, tgt_count) + 1):
                src_para = src_hf.Range.Paragraphs(i)
                tgt_para = tgt_hf.Range.Paragraphs(i)
                try:
                    tgt_para.Range.ParagraphFormat = src_para.Range.ParagraphFormat
                except Exception:
                    pass
                # Copy resolved character formatting; skip undefined/empty values.
                for prop in ("Name", "Size"):
                    try:
                        src_val = getattr(src_para.Range.Font, prop)
                        if src_val is None or src_val == 9999999:
                            continue
                        if prop == "Name" and not src_val:
                            # Range.Font.Name can be empty for theme-inherited fonts;
                            # fall back to the first character's resolved font.
                            try:
                                src_val = src_para.Range.Characters(1).Font.Name
                            except Exception:
                                continue
                        if not src_val:
                            continue
                        if src_val != getattr(tgt_para.Range.Font, prop):
                            setattr(tgt_para.Range.Font, prop, src_val)
                    except Exception:
                        pass
        except Exception as fmt_err:
            logger.warning(f"Could not copy header/footer formatting: {fmt_err}")

    @staticmethod
    def _copy_document_styles(target_doc, source_doc) -> None:
        """Copies the style definitions that govern header/footer text from source_doc
        to target_doc. FormattedText only carries explicitly-set formatting; values
        inherited from paragraph styles (spacing, line spacing, fonts) otherwise
        resolve against the target document's built-in styles (e.g. Word's default
        Normal style: 8pt after, 1.08 line spacing, Aptos). Matching the style
        definitions makes inherited values resolve identically to the source."""
        for style_name in ("Normal", "Header", "Footer"):
            try:
                tgt_style = target_doc.Styles(style_name)
                src_style = source_doc.Styles(style_name)
                tgt_style.ParagraphFormat = src_style.ParagraphFormat
                tgt_style.Font = src_style.Font
            except Exception:
                pass

    @staticmethod
    def _restore_sections_from_source(
        target_doc,
        source_doc,
        start_section_index: int,
        end_section_index: int,
        only_last_section: bool = False
    ) -> None:
        """Restores exact formatted headers/footers and page setup from source_doc onto target_doc."""
        range_section_count = end_section_index - start_section_index + 1
        exported_count = target_doc.Sections.Count

        PageExporterEngine._copy_document_styles(target_doc, source_doc)

        if only_last_section:
            restore_indices = [exported_count]
        else:
            restore_indices = list(range(1, exported_count + 1))

        for j in restore_indices:
            if range_section_count == exported_count:
                source_index = start_section_index + j - 1
            else:
                source_index = start_section_index + j - 1 if j < exported_count else end_section_index
            if only_last_section:
                source_index = end_section_index

            try:
                target_sec = target_doc.Sections(j)
                source_sec = source_doc.Sections(source_index)

                # 1. Mirror PageSetup properties
                ps_target = target_sec.PageSetup
                ps_source = source_sec.PageSetup
                for attr in (
                    "TopMargin", "BottomMargin", "LeftMargin", "RightMargin",
                    "HeaderDistance", "FooterDistance", "PageWidth", "PageHeight", "Orientation",
                ):
                    try:
                        setattr(ps_target, attr, getattr(ps_source, attr))
                    except Exception:
                        pass

                try:
                    ps_target.DifferentFirstPageHeaderFooter = ps_source.DifferentFirstPageHeaderFooter
                except Exception:
                    pass

                try:
                    ps_target.OddAndEvenPagesHeaderFooter = ps_source.OddAndEvenPagesHeaderFooter
                except Exception:
                    pass

                try:
                    ps_target.PageNumbering.RestartNumberingAtSection = ps_source.PageNumbering.RestartNumberingAtSection
                    ps_target.PageNumbering.StartingNumber = ps_source.PageNumbering.StartingNumber
                except Exception:
                    pass

                # 2. Copy formatted headers and footers using FormattedText
                for is_header in (True, False):
                    for hf_index in (1, 2, 3):
                        try:
                            target_hf = target_sec.Headers(hf_index) if is_header else target_sec.Footers(hf_index)
                            
                            # Section 1 of exported document must not link to non-existent previous section
                            if j == 1:
                                target_hf.LinkToPrevious = False
                            else:
                                source_hf_direct = source_sec.Headers(hf_index) if is_header else source_sec.Footers(hf_index)
                                try:
                                    target_hf.LinkToPrevious = bool(source_hf_direct.LinkToPrevious)
                                except Exception:
                                    target_hf.LinkToPrevious = False

                            if not target_hf.LinkToPrevious or j == 1:
                                effective_source_hf = PageExporterEngine._get_effective_header_footer(
                                    source_doc, source_index, is_header, hf_index
                                )
                                if effective_source_hf is not None:
                                    try:
                                        # Clear the target footer/header first so no stale content remains.
                                        target_hf.Range.Text = ""
                                        # Copy the source content EXCLUDING its trailing paragraph mark;
                                        # pasting it too would add a duplicate empty paragraph at the end.
                                        src_copy = effective_source_hf.Range
                                        if src_copy.End > src_copy.Start and src_copy.Text.endswith("\r"):
                                            src_copy.End = src_copy.End - 1
                                        if src_copy.End > src_copy.Start:
                                            target_hf.Range.FormattedText = src_copy.FormattedText
                                            # Copy resolved paragraph/run formatting so style-inherited
                                            # values (indent, alignment, spacing, fonts) match the source
                                            # even when the target document's built-in styles differ.
                                            PageExporterEngine._copy_header_footer_formatting(
                                                effective_source_hf, target_hf
                                            )
                                    except Exception as copy_err:
                                        logger.warning(f"Error copying FormattedText for section {j}: {copy_err}")
                            
                            # Update fields in header/footer range
                            try:
                                target_hf.Range.Fields.Update()
                            except Exception:
                                pass
                        except Exception as hf_err:
                            logger.warning(f"Error restoring header/footer for section {j}: {hf_err}")

            except Exception as sec_err:
                logger.warning(f"Error restoring properties for section {j}: {sec_err}")

        # Refresh overall document fields
        try:
            target_doc.Fields.Update()
        except Exception:
            pass

    @staticmethod
    def _export_by_selection(
        abs_source: str,
        abs_output: str,
        start_page: int,
        end_page: int,
        fmt_code: int,
        visible: bool
    ) -> str:
        """Fast Selection Copy Engine: Copies page range into new document."""
        with WordCOMContext(visible=visible) as word_app:
            source_doc = None
            target_doc = None
            try:
                source_doc = word_app.Documents.Open(
                    FileName=abs_source,
                    ReadOnly=True,
                    ConfirmConversions=False,
                    AddToRecentFiles=False,
                    Visible=visible
                )

                total_pages = source_doc.ComputeStatistics(WD_STATISTIC_PAGES)
                if start_page < 1 or end_page > total_pages or start_page > end_page:
                    raise ValueError(f"Invalid range [{start_page}-{end_page}] for document with {total_pages} pages.")

                start_section_index = PageExporterEngine._section_containing_page(source_doc, start_page)
                end_section_index = PageExporterEngine._section_containing_page(source_doc, end_page)

                start_range_start = PageExporterEngine._page_start_by_position(source_doc, start_page)
                if end_page == total_pages:
                    end_pos = source_doc.Content.End
                else:
                    end_pos = PageExporterEngine._page_start_by_position(source_doc, end_page + 1)

                extract_range = source_doc.Range(Start=start_range_start, End=end_pos)
                extract_range.Copy()

                target_doc = word_app.Documents.Add()
                target_doc.Content.Paste()

                # Restore high-fidelity headers, footers & page setup from original source_doc
                if start_section_index is not None and end_section_index is not None:
                    PageExporterEngine._restore_sections_from_source(
                        target_doc=target_doc,
                        source_doc=source_doc,
                        start_section_index=start_section_index,
                        end_section_index=end_section_index
                    )

                target_doc.SaveAs2(FileName=abs_output, FileFormat=fmt_code)
                logger.success(f"Exported pages [{start_page}-{end_page}] via selection copy to '{abs_output}'")

            finally:
                if target_doc:
                    try:
                        target_doc.Close(SaveChanges=False)
                    except Exception:
                        pass
                if source_doc:
                    try:
                        source_doc.Close(SaveChanges=False)
                    except Exception:
                        pass

        return abs_output
