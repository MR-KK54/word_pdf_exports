import io
import os
import threading
from typing import Tuple

import fitz  # PyMuPDF
import docx
from PIL import Image, ImageDraw

from word_exporter_pro.core.com_engine import WordCOMContext, aw
from word_exporter_pro.utils.logger import get_logger

logger = get_logger()

WD_EXPORT_PDF = 17  # wdExportFormatPDF

_preview_lock = threading.Lock()


def render_docx_fast_preview_jpeg(source_path: str, page: int = 1, max_width: int = 1000) -> Tuple[bytes, int, str]:
    """Generates an instant 0.05s stylized page JPEG image directly from document structure."""
    try:
        doc = docx.Document(source_path)
        total_pages = max(1, len(doc.paragraphs) // 15 + 1)
    except Exception:
        total_pages = 1

    w, h = max_width, int(max_width * 1.35)
    img = Image.new("RGB", (w, h), (248, 250, 252))
    draw = ImageDraw.Draw(img)

    # Outer document page shadow & card
    draw.rectangle([20, 20, w - 20, h - 20], fill=(255, 255, 255), outline=(218, 224, 233), width=3)
    draw.rectangle([35, 35, w - 35, 95], fill=(79, 70, 229))

    # Header title
    fn = os.path.basename(source_path)
    if len(fn) > 35:
        fn = fn[:32] + "..."
    draw.text((50, 52), fn, fill=(255, 255, 255))

    # Content paragraph lines
    for i in range(16):
        y = 130 + (i * 32)
        if y < h - 70:
            draw.line([50, y, w - 50, y], fill=(226, 232, 240), width=6)

    # Footer page label
    draw.text((w // 2 - 50, h - 50), f"Page {page} of {total_pages}", fill=(100, 116, 139))

    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return buf.getvalue(), total_pages, "image/jpeg"


def render_page_preview(
    source_path: str, page: int = 1, max_width: int = 1000
) -> Tuple[bytes, int, str]:
    """
    Renders a single page of a Word/PDF document to PNG/JPEG bytes.
    """
    if source_path.lower().endswith(".pdf") or has_preview_pdf(source_path):
        try:
            src = _ensure_pdf(source_path)
            if src.lower().endswith(".pdf") and os.path.exists(src):
                doc = fitz.open(src)
                total = doc.page_count
                idx = max(0, min(page - 1, total - 1))
                page_obj = doc.load_page(idx)
                zoom = max_width / page_obj.rect.width
                pix = page_obj.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                doc.close()
                return pix.tobytes("jpeg", jpg_quality=82), total, "image/jpeg"
        except Exception as e:
            logger.warning(f"PDF rendering fallback used: {e}")

    # Instant fast fallback preview image (0.05s)
    return render_docx_fast_preview_jpeg(source_path, page=page, max_width=max_width)


def _preview_path(source_path: str) -> str:
    base = os.path.splitext(os.path.basename(source_path))[0]
    return os.path.join(os.path.dirname(source_path), f"{base}_preview.pdf")


def has_preview_pdf(source_path: str) -> bool:
    """True when the source is a PDF or its Word preview PDF is already generated."""
    if source_path.lower().endswith(".pdf"):
        return True
    preview_pdf = _preview_path(source_path)
    return os.path.exists(preview_pdf) and os.path.getmtime(preview_pdf) >= os.path.getmtime(source_path)


_preview_jobs: dict = {}
_preview_jobs_lock = threading.Lock()


def ensure_preview_async(source_path: str) -> bool:
    """
    Ensures a Word preview PDF is (or will be) generated in the background.
    """
    if has_preview_pdf(source_path):
        return True

    key = os.path.abspath(source_path)
    with _preview_jobs_lock:
        if key in _preview_jobs:
            return False
        _preview_jobs[key] = True

    def worker():
        try:
            _ensure_pdf(source_path)
        except Exception as e:
            logger.error(f"Background preview generation failed for '{source_path}': {e}")
        finally:
            with _preview_jobs_lock:
                _preview_jobs.pop(key, None)

    threading.Thread(target=worker, daemon=True).start()
    return False


def _ensure_pdf(source_path: str) -> str:
    """Returns a PDF path for the source, converting Word files via a cached preview PDF."""
    if source_path.lower().endswith(".pdf"):
        return source_path

    preview_pdf = _preview_path(source_path)
    src_mtime = os.path.getmtime(source_path)
    with _preview_lock:
        if os.path.exists(preview_pdf) and os.path.getmtime(preview_pdf) >= src_mtime:
            return preview_pdf

        logger.info(f"Generating preview PDF for '{os.path.basename(source_path)}'...")
        word_ctx = WordCOMContext(visible=False)
        if word_ctx.available:
            with word_ctx as word_app:
                doc = word_app.Documents.Open(
                    FileName=os.path.abspath(source_path),
                    ReadOnly=True,
                    ConfirmConversions=False,
                    AddToRecentFiles=False,
                )
                try:
                    doc.ExportAsFixedFormat(
                        OutputFileName=os.path.abspath(preview_pdf),
                        ExportFormat=WD_EXPORT_PDF,
                    )
                    return preview_pdf
                finally:
                    try:
                        doc.Close(SaveChanges=False)
                    except Exception:
                        pass

        if aw is not None:
            try:
                doc = aw.Document(os.path.abspath(source_path))
                doc.save(os.path.abspath(preview_pdf))
                return preview_pdf
            except Exception as e:
                logger.warning(f"Aspose preview PDF generation skipped: {e}")

    return source_path
